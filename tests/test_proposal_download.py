import asyncio
import io
import json
import sqlite3

from docx import Document

import web_app
from web_app import _build_proposal_docx, api_profile_proposal_download


def _run(coro):
    return asyncio.get_event_loop().run_until_complete(coro)


def _body(response):
    return json.loads(response.body)


class _FakeRequest:
    def __init__(self, body=None, cookies=None):
        self._body = body or {}
        self.cookies = cookies or {}

    async def json(self):
        return self._body


def _paragraph_texts(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def test_build_proposal_docx_includes_all_populated_sections():
    proposal = {
        "background": "AI ethics in imagery is under-examined.",
        "objectives": "Understand consent and bias issues.",
        "research_questions": "- How does AI challenge consent?\n- How does bias manifest in outputs?",
        "related_work": "- Manovich, Artificial Aesthetics\n- Paglen, image ethics",
        "methodology": "- Historical analysis\n- Case studies",
        "expected_outcomes": "- A framework for ethical guidelines\n- A set of recommendations",
    }
    docx_bytes = _build_proposal_docx("Jane Doe", proposal)
    texts = _paragraph_texts(docx_bytes)

    assert any("Jane Doe" in t for t in texts)
    assert any("Introduction" in t or "Background" in t for t in texts)
    assert "AI ethics in imagery is under-examined." in texts
    assert "How does AI challenge consent?" in texts
    assert "Manovich, Artificial Aesthetics" in texts
    assert "Historical analysis" in texts
    assert "A framework for ethical guidelines" in texts


def test_build_proposal_docx_skips_empty_sections():
    proposal = {
        "background": "Some background.",
        "objectives": "Some objectives.",
        "research_questions": "A single research question.",
        "related_work": "",
        "methodology": "Some methodology.",
        "expected_outcomes": "",
    }
    docx_bytes = _build_proposal_docx("Jane Doe", proposal)
    texts = _paragraph_texts(docx_bytes)

    assert not any("Relevant Literature" in t for t in texts)
    assert not any("Expected Outcomes" in t for t in texts)


def test_build_proposal_docx_renders_bullet_lines_as_list_items():
    proposal = {
        "background": "bg", "objectives": "obj",
        "research_questions": "- First question\n- Second question",
        "related_work": "", "methodology": "single method line", "expected_outcomes": "",
    }
    docx_bytes = _build_proposal_docx("Jane Doe", proposal)
    doc = Document(io.BytesIO(docx_bytes))

    bullet_paragraphs = [p for p in doc.paragraphs if p.text in ("First question", "Second question")]
    assert len(bullet_paragraphs) == 2
    for p in bullet_paragraphs:
        assert p.style.name == "List Bullet"

    plain_paragraphs = [p for p in doc.paragraphs if p.text == "single method line"]
    assert len(plain_paragraphs) == 1
    assert plain_paragraphs[0].style.name != "List Bullet"


def _init_db(tmp_path, monkeypatch):
    db_path = tmp_path / "test_faculty.db"
    monkeypatch.setattr(web_app, "DB_PATH", str(db_path))
    web_app._init_profiles_db()
    web_app._auth_sessions.clear()
    return db_path


def _signup_session(email="jane@depaul.edu"):
    _run(web_app.api_auth_signup(_FakeRequest({"email": email, "password": "hunter222"})))
    return list(web_app._auth_sessions.keys())[-1]


def test_download_requires_login(tmp_path, monkeypatch):
    _init_db(tmp_path, monkeypatch)
    response = _run(api_profile_proposal_download(_FakeRequest(cookies={})))
    assert response.status_code == 401


def test_download_returns_404_when_no_profile(tmp_path, monkeypatch):
    _init_db(tmp_path, monkeypatch)
    token = _signup_session()
    response = _run(api_profile_proposal_download(_FakeRequest(cookies={"session_token": token})))
    assert response.status_code == 404


def test_download_returns_404_when_no_proposal(tmp_path, monkeypatch):
    db_path = _init_db(tmp_path, monkeypatch)
    token = _signup_session()
    user_id = web_app._auth_sessions[token]
    con = sqlite3.connect(db_path)
    con.execute("INSERT INTO profiles (name, user_id) VALUES ('Jane Doe', ?)", (user_id,))
    con.commit()
    con.close()

    response = _run(api_profile_proposal_download(_FakeRequest(cookies={"session_token": token})))
    assert response.status_code == 404


def test_download_returns_docx_when_proposal_exists(tmp_path, monkeypatch):
    db_path = _init_db(tmp_path, monkeypatch)
    token = _signup_session()
    user_id = web_app._auth_sessions[token]
    con = sqlite3.connect(db_path)
    cur = con.execute("INSERT INTO profiles (name, user_id) VALUES ('Jane Doe', ?)", (user_id,))
    profile_id = cur.lastrowid
    con.execute(
        "INSERT INTO proposals (profile_id, background, objectives, research_questions, methodology) "
        "VALUES (?, 'bg text', 'obj text', 'rq text', 'method text')", (profile_id,)
    )
    con.commit()
    con.close()

    response = _run(api_profile_proposal_download(_FakeRequest(cookies={"session_token": token})))
    assert response.status_code == 200
    assert response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "attachment" in response.headers["content-disposition"]
    assert "Jane_Doe" in response.headers["content-disposition"]

    texts = _paragraph_texts(response.body)
    assert "bg text" in texts
